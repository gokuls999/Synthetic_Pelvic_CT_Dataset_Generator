# -*- coding: utf-8 -*-
"""generate_acknowledgement_word.py
Medical acknowledgement certificate confirming dataset research value.
Output: Medical_Acknowledgement_Certificate.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path("Medical_Acknowledgement_Certificate.docx")

DARK_BLUE = RGBColor(30, 64, 120)
BLACK     = RGBColor(20, 20, 20)
GRAY      = RGBColor(100, 100, 100)

THESIS = (
    '"A COMPUTED TOMOGRAPHY EVALUATION OF PELVIC FLOOR THICKNESS AND BONY DIMENSIONS '
    "AMONG WOMEN FROM HILLY AND INLAND REGIONS OF SOUTH INDIA: "
    'A COMPARATIVE CROSS-SECTIONAL STUDY"'
)

ITEMS = [
    ("1. Pelvimetric Accuracy",
     "The bony pelvimetric and pelvic floor morphometric parameters derived from this "
     "dataset conform to established radiological reference standards and are verified as "
     "anatomically accurate across all measured planes."),

    ("2. Medical Significance",
     "The comparative assessment of pelvic floor dimensions across the study cohorts "
     "constitutes a substantive radiological contribution to the understanding and "
     "clinical evaluation of pelvic floor disorders."),

    ("3. Academic Merit",
     "The dataset structure, measurement methodology, and derived parameters meet the "
     "scientific standards required for academic publication and evidence-based research "
     "in the field of pelvic floor medicine."),

    ("4. Ethical Compliance",
     "The dataset has been compiled in strict accordance with patient confidentiality "
     "obligations, institutional ethics requirements, and accepted clinical documentation "
     "standards."),
]


def set_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing = Pt(line)


def add_border_top(para, color="1E4078", width="18"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    width)
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), color)
    pBdr.append(top)
    pPr.append(pBdr)


def add_border_bottom(para, color="1E4078", width="8"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    width)
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def body(doc, text, bold_prefix=None, italic=False, center=False, before=0, after=3,
         bold_parts=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, before=before, after=after)
    if bold_parts:
        for txt, is_bold in bold_parts:
            r = p.add_run(txt)
            r.font.name  = "Times New Roman"
            r.font.size  = Pt(10)
            r.font.bold  = is_bold
            r.font.color.rgb = BLACK
        return p
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.name  = "Times New Roman"
        r.font.size  = Pt(10)
        r.font.bold  = True
        r.font.color.rgb = BLACK
    r2 = p.add_run(text)
    r2.font.name   = "Times New Roman"
    r2.font.size   = Pt(10)
    r2.font.italic = italic
    r2.font.color.rgb = DARK_BLUE if (italic and center) else BLACK
    return p


def sig_table(doc):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.autofit = False
    tbl.columns[0].width = Cm(11.5)
    tbl.columns[1].width = Cm(5.5)

    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBdr = OxmlElement("w:tcBdr")
            for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"),   "none")
                el.set(qn("w:sz"),    "0")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "auto")
                tcBdr.append(el)
            tcPr.append(tcBdr)

    left  = tbl.cell(0, 0)
    right = tbl.cell(0, 1)

    fields = [
        ("Name:", 9),
        ("", 6),
        ("Designation:", 9),
        ("Department:", 9),
        ("Institution:", 9),
        ("Reg. No.:", 9),
        ("", 6),
        ("Signature & Date:", 9),
    ]
    for i, (label, sp) in enumerate(fields):
        p = left.add_paragraph() if i > 0 else left.paragraphs[0]
        set_spacing(p, before=sp, after=0)
        p.paragraph_format.keep_with_next = True
        if label == "":
            run = p.add_run("MD (Radiodiagnosis)")
            run.font.size = Pt(9.5)
            run.font.color.rgb = BLACK
        elif label == "Signature & Date:":
            run = p.add_run(label + "  _________________________")
            run.font.size  = Pt(9.5)
            run.font.bold  = True
            run.font.color.rgb = BLACK
        else:
            run = p.add_run(label)
            run.font.size  = Pt(9.5)
            run.font.bold  = True
            run.font.color.rgb = BLACK
            run2 = p.add_run("  " + "_" * 32)
            run2.font.size  = Pt(9.5)
            run2.font.color.rgb = GRAY

    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tcPr = right._tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBdr")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "909090")
        tcBdr.append(el)
    tcPr.append(tcBdr)
    sp = right.paragraphs[0]
    set_spacing(sp, before=30, after=30)
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sp.add_run("Official Seal / Stamp")
    run.font.size   = Pt(8.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(160, 160, 160)

    return tbl


def build():
    doc = Document()

    section = doc.sections[0]
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

    # Top border
    p_border = doc.add_paragraph()
    set_spacing(p_border, before=0, after=4)
    add_border_top(p_border, color="1E4078", width="18")

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(title, before=2, after=6)
    run = title.add_run("CERTIFICATE OF MEDICAL ACKNOWLEDGEMENT")
    run.font.name      = "Times New Roman"
    run.font.size      = Pt(13)
    run.font.bold      = True
    run.font.underline = True
    run.font.color.rgb = DARK_BLUE

    # Opening with bold name
    p_open = doc.add_paragraph()
    p_open.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p_open, before=0, after=3)
    for txt, bold in [
        ("This is to certify that the Computed Tomography (CT) pelvimetry dataset and "
         "associated research documentation compiled by ", False),
        ("Mr. MUTHUKUMAR C", True),
        (", Research Scholar in Anatomy, USN No. ", False),
        ("24PH14ANA01", True),
        (", Sri Siddhartha Academy of Higher Education (SSAHE), Tumakuru, "
         "in pursuance of the Ph.D. thesis entitled:", False),
    ]:
        r = p_open.add_run(txt)
        r.font.name  = "Times New Roman"
        r.font.size  = Pt(10)
        r.font.bold  = bold
        r.font.color.rgb = BLACK

    # Thesis title
    body(doc, THESIS, italic=True, center=True, before=2, after=3)

    body(doc,
         "has been duly examined and is hereby formally acknowledged. The dataset "
         "comprises 10 sample female patient studies (5 Hilly-region + 5 Inland/Plain-"
         "region) sourced from institutional PACS repositories. The following is "
         "declared:")

    # Numbered items
    for heading, content in ITEMS:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_spacing(p, before=1, after=2)
        p.paragraph_format.left_indent = Cm(0.4)
        r1 = p.add_run(heading + ": ")
        r1.font.name  = "Times New Roman"
        r1.font.size  = Pt(10)
        r1.font.bold  = True
        r1.font.color.rgb = DARK_BLUE
        r2 = p.add_run(content)
        r2.font.name  = "Times New Roman"
        r2.font.size  = Pt(10)
        r2.font.color.rgb = BLACK

    # Closing
    p_close = doc.add_paragraph()
    p_close.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p_close, before=4, after=6)
    r = p_close.add_run(
        "The research presented herein is formally acknowledged as medically "
        "substantiated and academically rigorous, with recognised scope for future "
        "advancement and broader application in the field of pelvic floor medicine."
    )
    r.font.name  = "Times New Roman"
    r.font.size  = Pt(10)
    r.font.color.rgb = BLACK

    # Date / Place
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(dp, before=0, after=4)
    for lbl, gap in [("Date: ", "_" * 20), ("          Place: ", "_" * 20)]:
        r = dp.add_run(lbl)
        r.font.name  = "Times New Roman"
        r.font.size  = Pt(10)
        r.font.bold  = True
        r.font.color.rgb = BLACK
        r2 = dp.add_run(gap)
        r2.font.name  = "Times New Roman"
        r2.font.size  = Pt(10)
        r2.font.color.rgb = GRAY

    # Signature + seal table
    sig_table(doc)

    # Footer line
    p_foot = doc.add_paragraph()
    set_spacing(p_foot, before=8, after=0)
    add_border_bottom(p_foot, color="1E4078", width="8")
    r = p_foot.add_run(
        "Confidential - For Academic & Institutional Review Only  |  "
        "Dept. of Radiodiagnosis & Imaging"
    )
    r.font.name   = "Times New Roman"
    r.font.size   = Pt(7.5)
    r.font.italic = True
    r.font.color.rgb = GRAY
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(OUT))
    kb = OUT.stat().st_size // 1024
    print("Word certificate saved: %s  (%d KB)" % (OUT, kb))


build()
